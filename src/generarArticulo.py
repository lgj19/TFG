from docx import Document
from entities.Equipo import Equipo
from recuEntitiesBD import *


dictEqG = {}
dictEqP = {}
dictPartido = {}
dictMVP1G = {}
dictMVP2G = {}
dictMVP1P = {}
dictMVP2P = {}



def inyeccionPlantilla(nombreEqL, nombreEqV, temporada, tiposParrafos, plantilla):
    #Recuperar equipos
    eqL = recuEqNombre(nombreEqL, temporada)
    eqV = recuEqNombre(nombreEqV, temporada)
    print(eqL, eqV)

    #Recuperamos el partido
    partido = recuPartido(eqL.id, eqV.id, temporada)
    print(partido)
    #Recuperamos los jugadores de los equipos
    jugsEqL = recuJugsEqL(eqL.id, partido.id)
    jugsEqV = recuJugsEqV(eqV.id, partido.id)

    #Rellenar los diccionarios
    rellenarDicts(eqL, eqV, partido, jugsEqL, jugsEqV)
    
    return inyeccionParrafos(plantilla, tiposParrafos)



def inyeccionParrafos(plantilla, tiposParrafos):
    global dictEqG, dictEqP, dictPartido, dictMVP1G, dictMVP2G, dictMVP1P, dictMVP2P
    dicts = [dictEqG, dictEqP, dictPartido,dictMVP1G, dictMVP2G, dictMVP1P, dictMVP2P]
    
    articuloFinal = Document()
    newParrafo = ""

    for parrafo in plantilla.parrafos:
        if parrafo.tipoParrafo in tiposParrafos:
            articuloFinal.add_heading(parrafo.tipoParrafo, level=2)
            newParrafo = parrafo.contenido
            for d in dicts:
                for key, val in d.items():
                    newParrafo = newParrafo.replace(key, str(val))
            articuloFinal.add_paragraph(newParrafo)

    return articuloFinal

        


def rellenarDicts(eqL, eqV, partido, jugsEqL, jugsEqV):
    eqG = []; eqP = []; 
    global dictEqG, dictEqP, dictPartido, dictMVP1G, dictMVP2G, dictMVP1P, dictMVP2P

    #Asignamos los equipos ganadores y perdedores a partir de los puntos
    #del equipo local y visitante
    if(partido.puntosTotalLocal > partido.puntosTotalVisitante):
        eqG = eqL
        eqP = eqV
        dictEqG['LOCAL'] = True; dictEqP['LOCAL'] = False
    else:
        eqG = eqV
        eqP = eqL
        dictEqG['LOCAL'] = False; dictEqP['LOCAL'] = True

    #Introducimos los datos para los equipos ganadores y perdedores
    dictEqG["NOMBRE_G"] = eqG.nombre; dictEqP["NOMBRE_P"] = eqP.nombre
    dictEqG["SIGLAS_G"] = eqG.siglas; dictEqP["SIGLAS_P"] = eqP.siglas
    dictEqG["FUND_G"] = eqG.anyoFundacion; dictEqP["FUND_P"] = eqP.anyoFundacion
    dictEqG["TEMP_G"] = eqG.temporada; dictEqP["TEMP_P"] = eqP.temporada

    #Introducimos los datos genéricos del partido
    dictPartido["FASE_P"] = partido.faseCompeticion
    dictPartido["JORNADA_P"] = partido.jornada
    dictPartido["PABELLON"] = partido.pabellon
    dictPartido["PUBLICO"] = partido.publico

    #Introducimos los puntos de cada equipo
    if(dictEqG['LOCAL']):
        dictPartido["PUNTOS_G"] = partido.puntosTotalLocal
        dictPartido["PUNTOS1Q_G"] = partido.puntos1QLocal
        dictPartido["PUNTOS2Q_G"] = partido.puntos2QLocal
        dictPartido["PUNTOS3Q_G"] = partido.puntos3QLocal
        dictPartido["PUNTOS4Q_G"] = partido.puntos4QLocal
        dictPartido["PUNTOSEQ_G"] = partido.puntosExtraLocal
        dictPartido["PUNTOS_P"] = partido.puntosTotalVisitante
        dictPartido["PUNTOS1Q_P"] = partido.puntos1QVisitante
        dictPartido["PUNTOS2Q_P"] = partido.puntos2QVisitante
        dictPartido["PUNTOS3Q_P"] = partido.puntos3QVisitante
        dictPartido["PUNTOS4Q_P"] = partido.puntos4QVisitante
        dictPartido["PUNTOSEQ_P"] = partido.puntosExtraVisitante
    else:
        dictPartido["PUNTOS_P"] = partido.puntosTotalLocal
        dictPartido["PUNTOS1Q_P"] = partido.puntos1QLocal
        dictPartido["PUNTOS2Q_P"] = partido.puntos2QLocal
        dictPartido["PUNTOS3Q_P"] = partido.puntos3QLocal
        dictPartido["PUNTOS4Q_P"] = partido.puntos4QLocal
        dictPartido["PUNTOSEQ_P"] = partido.puntosExtraLocal
        dictPartido["PUNTOS_G"] = partido.puntosTotalVisitante
        dictPartido["PUNTOS1Q_G"] = partido.puntos1QVisitante
        dictPartido["PUNTOS2Q_G"] = partido.puntos2QVisitante
        dictPartido["PUNTOS3Q_G"] = partido.puntos3QVisitante
        dictPartido["PUNTOS4Q_G"] = partido.puntos4QVisitante
        dictPartido["PUNTOSEQ_G"] = partido.puntosExtraVisitante
    dictPartido["PUNTOS_L"] = partido.puntosTotalLocal
    dictPartido["PUNTOS1Q_L"] = partido.puntos1QLocal
    dictPartido["PUNTOS2Q_L"] = partido.puntos2QLocal
    dictPartido["PUNTOS3Q_L"] = partido.puntos3QLocal
    dictPartido["PUNTOS4Q_L"] = partido.puntos4QLocal
    dictPartido["PUNTOSEQ_L"] = partido.puntosExtraLocal
    dictPartido["PUNTOS_V"] = partido.puntosTotalVisitante
    dictPartido["PUNTOS1Q_V"] = partido.puntos1QVisitante
    dictPartido["PUNTOS2Q_V"] = partido.puntos2QVisitante
    dictPartido["PUNTOS3Q_V"] = partido.puntos3QVisitante
    dictPartido["PUNTOS4Q_V"] = partido.puntos4QVisitante
    dictPartido["PUNTOSEQ_V"] = partido.puntosExtraVisitante
    
    dictPartido["PUNTOS_1aPARTE_L"] = partido.puntos1QLocal + partido.puntos2QLocal
    dictPartido["PUNTOS_1aPARTE_V"] = partido.puntos1QVisitante + partido.puntos2QVisitante
    dictPartido["DIF_PUNTOS1Q"] = abs(partido.puntos1QLocal - partido.puntos1QVisitante)
    dictPartido["DIF_PUNTOS2Q"] = abs(partido.puntos2QLocal - partido.puntos2QVisitante)
    dictPartido["DIF_PUNTOS3Q"] = abs(partido.puntos3QLocal - partido.puntos3QVisitante)
    dictPartido["DIF_PUNTOS4Q"] = abs(partido.puntos4QLocal - partido.puntos4QVisitante)
    dictPartido["DIF_PUNTOSEQ"] = abs(partido.puntosExtraLocal - partido.puntosExtraVisitante)
    dictPartido["DIF_PUNTOS_1aPARTE"] = abs((partido.puntos1QLocal + partido.puntos2QLocal) - 
        (partido.puntos1QVisitante + partido.puntos2QVisitante))
    dictPartido["DIF_PUNTOS"] = abs(partido.puntosTotalLocal - partido.puntosTotalVisitante)

    if(partido.puntos1QLocal - partido.puntos1QVisitante > 0):
        if(dictEqG['LOCAL']):
            dictPartido["NOMBRE_1Q_G"] = eqG.nombre
        else:
            dictPartido["NOMBRE_1Q_G"] = eqP.nombre
    else:
        if(dictEqG['LOCAL']):
            dictPartido["NOMBRE_1Q_G"] = eqP.nombre
        else:
            dictPartido["NOMBRE_1Q_G"] = eqG.nombre

    if(partido.puntos2QLocal - partido.puntos2QVisitante > 0):
        if(dictEqG['LOCAL']):
            dictPartido["NOMBRE_2Q_G"] = eqG.nombre
        else:
            dictPartido["NOMBRE_2Q_G"] = eqP.nombre
    else:
        if(dictEqG['LOCAL']):
            dictPartido["NOMBRE_2Q_G"] = eqP.nombre
        else:
            dictPartido["NOMBRE_2Q_G"] = eqG.nombre

    if(partido.puntos3QLocal - partido.puntos3QVisitante > 0):
        if(dictEqG['LOCAL']):
            dictPartido["NOMBRE_3Q_G"] = eqG.nombre
        else:
            dictPartido["NOMBRE_3Q_G"] = eqP.nombre
    else:
        if(dictEqG['LOCAL']):
            dictPartido["NOMBRE_3Q_G"] = eqP.nombre
        else:
            dictPartido["NOMBRE_3Q_G"] = eqG.nombre

    if(partido.puntos4QLocal - partido.puntos4QVisitante > 0):
        if(dictEqG['LOCAL']):
            dictPartido["NOMBRE_4Q_G"] = eqG.nombre
        else:
            dictPartido["NOMBRE_4Q_G"] = eqP.nombre
    else:
        if(dictEqG['LOCAL']):
            dictPartido["NOMBRE_4Q_G"] = eqP.nombre
        else:
            dictPartido["NOMBRE_4Q_G"] = eqG.nombre

    if(partido.puntosExtraLocal - partido.puntosExtraVisitante > 0):
        if(dictEqG['LOCAL']):
            dictPartido["NOMBRE_EQ_G"] = eqG.nombre
        else:
            dictPartido["NOMBRE_EQ_G"] = eqP.nombre
    else:
        if(dictEqG['LOCAL']):
            dictPartido["NOMBRE_EQ_G"] = eqP.nombre
        else:
            dictPartido["NOMBRE_EQ_G"] = eqG.nombre
    

    

    #Recuperamos los 2 mejores jugadores del partido de cada equipo
    jugMVP1G = jugMVP2G = jugMVP1P = jugMVP2P = []

    jugMVP1G = jugsEqL[0]; jugMVP2G = jugsEqL[1]; jugMVP1P = jugsEqV[0]; jugMVP2P = jugsEqV[1]
    
    for jug in jugsEqL:
        if(jug.valoracion > jugMVP1G.valoracion):
            jugMVP2G = jugMVP1G
            jugMVP1G = jug
        elif(jug.valoracion > jugMVP2G.valoracion and jug != jugMVP1G):
            jugMVP2G = jug

    for jug in jugsEqV:
        if(jug.valoracion > jugMVP1P.valoracion):
            jugMVP2P = jugMVP1P
            jugMVP1P = jug
        elif(jug.valoracion > jugMVP2P.valoracion  and jug != jugMVP1P):
            jugMVP2P = jug

    #Invertimos valores en caso de que no haya ganado el equipo local
    if(dictEqG['LOCAL'] == False):
        aux = jugMVP1G
        jugMVP1G = jugMVP1P
        jugMVP1P = aux
        aux = jugMVP2G
        jugMVP2G = jugMVP2P
        jugMVP2P = aux

    #Introducimos los jugadores selecionados en los diccionarios
    dictMVP1G = {"JUG_G1_APODO":jugMVP1G.nombre, "JUG_G1_NOMBRE_COMPLETO":jugMVP1G.nombreCompleto,
      "JUG_G1_PRIMER_APELLIDO":jugMVP1G.primerApellido,
      "JUG_G1_SEGUNDO_APELLIDO":jugMVP1G.segundoApellido,
      "JUG_G1_NACION_SIGLAS": jugMVP1G.nacionalidad, "JUG_G1_LUGAR_NAC":jugMVP1G.lugarNacimiento,
      "JUG_G1_FECHA_NAC":jugMVP1G.fechaNacimiento[:10], "JUG_G1_POSICION":jugMVP1G,
    "JUG_G1_ALTURA": jugMVP1G.altura, "JUG_G1_DEBUT_ACB":jugMVP1G.debutACB[:10], "JUG_G1_TWITTER":jugMVP1G.twitter, "JUG_G1_DORSAL": jugMVP1G.dorsal, "JUG_G1_MIN":jugMVP1G.minutos,
    "JUG_G1_PUNTOS":jugMVP1G.puntosTotales, "JUG_G1_T2":jugMVP1G.tirosDe2, "JUG_G1_T3":jugMVP1G.tirosDe3, "JUG_G1_T1":jugMVP1G.tirosDe1, "JUG_G1_INT_T2":jugMVP1G.intentosDe2,
    "JUG_G1_INT_T3":jugMVP1G.intentosDe3, "JUG_G1_INT_T1":jugMVP1G.intentosDe1, "JUG_G1_REB_DEF":jugMVP1G.rebotesDefensivos, "JUG_G1_REB_OF":jugMVP1G.rebotesOfensivos,
    "JUG_G1_ASIST":jugMVP1G.asistencias, "JUG_G1_ROBOS":jugMVP1G.robos, "JUG_G1_PERD":jugMVP1G.perdidas, "JUG_G1_CONTRAS":jugMVP1G.contraataques, "JUG_G1_TAP":jugMVP1G.tapones,
    "JUG_G1_TAP_RECIBIDOS":jugMVP1G.taponesRecibidos, "JUG_G1_MATES":jugMVP1G.mates, "JUG_G1_FALTAS":jugMVP1G.faltas, "JUG_G1_FALTAS_RECIBIDAS":jugMVP1G.faltasRecibidas,
    "JUG_G1_VALORACION":jugMVP1G.valoracion, "JUG_G1_MASMENOS":jugMVP1G.masMenos}
    
    dictMVP2G = {"JUG_G2_APODO":jugMVP2G.nombre, "JUG_G2_NOMBRE_COMPLETO":jugMVP2G.nombreCompleto, "JUG_G2_PRIMER_APELLIDO":jugMVP2G.primerApellido, "JUG_G2_SEGUNDO_APELLIDO":jugMVP2G.segundoApellido,
    "JUG_G2_NACION_SIGLAS": jugMVP2G.nacionalidad, "JUG_G2_LUGAR_NAC":jugMVP2G.lugarNacimiento, "JUG_G2_FECHA_NAC":jugMVP2G.fechaNacimiento[:10], "JUG_G2_POSICION":jugMVP2G,
    "JUG_G2_ALTURA": jugMVP2G.altura, "JUG_G2_DEBUT_ACB":jugMVP2G.debutACB[:10], "JUG_G2_TWITTER":jugMVP2G.twitter, "JUG_G2_DORSAL": jugMVP2G.dorsal, "JUG_G2_MIN":jugMVP2G.minutos,
    "JUG_G2_PUNTOS":jugMVP2G.puntosTotales, "JUG_G2_T2":jugMVP2G.tirosDe2, "JUG_G2_T3":jugMVP2G.tirosDe3, "JUG_G2_T1":jugMVP2G.tirosDe1, "JUG_G2_INT_T2":jugMVP2G.intentosDe2,
    "JUG_G2_INT_T3":jugMVP2G.intentosDe3, "JUG_G2_INT_T1":jugMVP2G.intentosDe1, "JUG_G2_REB_DEF":jugMVP2G.rebotesDefensivos, "JUG_G2_REB_OF":jugMVP2G.rebotesOfensivos,
    "JUG_G2_ASIST":jugMVP2G.asistencias, "JUG_G2_ROBOS":jugMVP2G.robos, "JUG_G2_PERD":jugMVP2G.perdidas, "JUG_G2_CONTRAS":jugMVP2G.contraataques, "JUG_G2_TAP":jugMVP2G.tapones,
    "JUG_G2_TAP_RECIBIDOS":jugMVP2G.taponesRecibidos, "JUG_G2_MATES":jugMVP2G.mates, "JUG_G2_FALTAS":jugMVP2G.faltas, "JUG_G2_FALTAS_RECIBIDAS":jugMVP2G.faltasRecibidas,
    "JUG_G2_VALORACION":jugMVP2G.valoracion, "JUG_G2_MASMENOS":jugMVP2G.masMenos}

    dictMVP1P = {"JUG_P1_APODO":jugMVP1P.nombre, "JUG_P1_NOMBRE_COMPLETO":jugMVP1P.nombreCompleto, "JUG_P1_PRIMER_APELLIDO":jugMVP1P.primerApellido, "JUG_P1_SEGUNDO_APELLIDO":jugMVP1P.segundoApellido,
    "JUG_P1_NACION_SIGLAS": jugMVP1P.nacionalidad, "JUG_P1_LUGAR_NAC":jugMVP1P.lugarNacimiento, "JUG_P1_FECHA_NAC":jugMVP1P.fechaNacimiento[:10], "JUG_P1_POSICION":jugMVP1P,
    "JUG_P1_ALTURA": jugMVP1P.altura, "JUG_P1_DEBUT_ACB":jugMVP1P.debutACB[:10], "JUG_P1_TWITTER":jugMVP1P.twitter, "JUG_P1_DORSAL": jugMVP1P.dorsal, "JUG_P1_MIN":jugMVP1P.minutos,
    "JUG_P1_PUNTOS":jugMVP1P.puntosTotales, "JUG_P1_T2":jugMVP1P.tirosDe2, "JUG_P1_T3":jugMVP1P.tirosDe3, "JUG_P1_T1":jugMVP1P.tirosDe1, "JUG_P1_INT_T2":jugMVP1P.intentosDe2,
    "JUG_P1_INT_T3":jugMVP1P.intentosDe3, "JUG_P1_INT_T1":jugMVP1P.intentosDe1, "JUG_P1_REB_DEF":jugMVP1P.rebotesDefensivos, "JUG_P1_REB_OF":jugMVP1P.rebotesOfensivos,
    "JUG_P1_ASIST":jugMVP1P.asistencias, "JUG_P1_ROBOS":jugMVP1P.robos, "JUG_P1_PERD":jugMVP1P.perdidas, "JUG_P1_CONTRAS":jugMVP1P.contraataques, "JUG_P1_TAP":jugMVP1P.tapones,
    "JUG_P1_TAP_RECIBIDOS":jugMVP1P.taponesRecibidos, "JUG_P1_MATES":jugMVP1P.mates, "JUG_P1_FALTAS":jugMVP1P.faltas, "JUG_P1_FALTAS_RECIBIDAS":jugMVP1P.faltasRecibidas,
    "JUG_P1_VALORACION":jugMVP1P.valoracion, "JUG_P1_MASMENOS":jugMVP1P.masMenos}

    dictMVP2P = {"JUG_P2_APODO":jugMVP2P.nombre, "JUG_P2_NOMBRE_COMPLETO":jugMVP2P.nombreCompleto, "JUG_P2_PRIMER_APELLIDO":jugMVP2P.primerApellido, "JUG_P2_SEGUNDO_APELLIDO":jugMVP2P.segundoApellido,
    "JUG_P2_NACION_SIGLAS": jugMVP2P.nacionalidad, "JUG_P2_LUGAR_NAC":jugMVP2P.lugarNacimiento, "JUG_P2_FECHA_NAC":jugMVP2P.fechaNacimiento[:10], "JUG_P2_POSICION":jugMVP2P,
    "JUG_P2_ALTURA": jugMVP2P.altura, "JUG_P2_DEBUT_ACB":jugMVP2P.debutACB[:10], "JUG_P2_TWITTER":jugMVP2P.twitter, "JUG_P2_DORSAL": jugMVP2P.dorsal, "JUG_P2_MIN":jugMVP2P.minutos,
    "JUG_P2_PUNTOS":jugMVP2P.puntosTotales, "JUG_P2_T2":jugMVP2P.tirosDe2, "JUG_P2_T3":jugMVP2P.tirosDe3, "JUG_P2_T1":jugMVP2P.tirosDe1, "JUG_P2_INT_T2":jugMVP2P.intentosDe2,
    "JUG_P2_INT_T3":jugMVP2P.intentosDe3, "JUG_P2_INT_T1":jugMVP2P.intentosDe1, "JUG_P2_REB_DEF":jugMVP2P.rebotesDefensivos, "JUG_P2_REB_OF":jugMVP2P.rebotesOfensivos,
    "JUG_P2_ASIST":jugMVP2P.asistencias, "JUG_P2_ROBOS":jugMVP2P.robos, "JUG_P2_PERD":jugMVP2P.perdidas, "JUG_P2_CONTRAS":jugMVP2P.contraataques, "JUG_P2_TAP":jugMVP2P.tapones,
    "JUG_P2_TAP_RECIBIDOS":jugMVP2P.taponesRecibidos, "JUG_P2_MATES":jugMVP2P.mates, "JUG_P2_FALTAS":jugMVP2P.faltas, "JUG_P2_FALTAS_RECIBIDAS":jugMVP2P.faltasRecibidas,
    "JUG_P2_VALORACION":jugMVP2P.valoracion, "JUG_P2_MASMENOS":jugMVP2P.masMenos}


    #Generar estadísticas globales de equipo ambos equipos
    jugsEqG = []; jugsEqP = []
    valoracion = 0; rebotesOf = 0; rebotesDef = 0; robos = 0; perdidas = 0; asistencias = 0
    tapones = 0; taponesRecibidos = 0; faltas = 0; faltasRecibidas = 0; contraataques = 0
    tirosDe1 = 0; tirosDe2 = 0; tirosDe3 = 0; intentosDe1 = 0; intentosDe2 = 0; intentosDe3 = 0
    porcentajeDe1 = 0; porcentajeDe2 = 0; porcentajeDe3 = 0

    if(dictEqG['LOCAL']):
        jugsEqG = jugsEqL; jugsEqP = jugsEqV
    else:
        jugsEqG = jugsEqV; jugsEqP = jugsEqL


    for jug in jugsEqG:
        valoracion += jug.valoracion
        rebotesDef += jug.rebotesDefensivos
        rebotesOf += jug.rebotesOfensivos
        robos += jug.robos
        perdidas += jug.perdidas
        asistencias += jug.asistencias
        tapones += jug.tapones
        taponesRecibidos += jug.taponesRecibidos
        faltas += jug.faltas
        faltasRecibidas += jug.faltasRecibidas
        contraataques += jug.contraataques
        tirosDe1 += jug.tirosDe1
        tirosDe2 += jug.tirosDe2
        tirosDe3 += jug.tirosDe3
        intentosDe1 += jug.intentosDe1
        intentosDe2 += jug.intentosDe2
        intentosDe3 += jug.intentosDe3
    porcentajeDe1 = tirosDe1 / intentosDe1 * 100
    porcentajeDe2 = tirosDe2 / intentosDe2 * 100
    porcentajeDe3 = tirosDe3 / intentosDe3 * 100

    dictEqG.update({"REBOTES_OF_G":rebotesOf, "REBOTES_DEF_G":rebotesDef,
      "FALTAS_G":faltas,
      "FALTAS_RECIBIDAS_G":faltasRecibidas, "TAPONES_G":tapones, 
      "TAPONES_RECIBIDOS_G":taponesRecibidos,
      "CONTRAATAQUES_G":contraataques, "ASISTENCIAS_G":asistencias,
      "PERDIDAS_G":perdidas,"VALORACION_G":valoracion,
      "ROBOS_G":robos, "TL_TIRADOS_G":intentosDe1,
      "TL_ANOTADOS_G":tirosDe1, "PORCENTAJE_TL_G":porcentajeDe1,
      "T2_TIRADOS_G":intentosDe2, "T2_ANOTADOS_G":tirosDe2, "PORCENTAJE_T2_G":porcentajeDe2,
      "T3_TIRADOS_G":intentosDe3, "T3_ANOTADOS_G":tirosDe3, "PORCENTAJE_T3_G":porcentajeDe3})
    
    valoracion = 0; rebotesOf = 0; rebotesDef = 0; robos = 0; perdidas = 0; asistencias = 0
    tapones = 0; taponesRecibidos = 0; faltas = 0; faltasRecibidas = 0; contraataques = 0
    tirosDe1 = 0; tirosDe2 = 0; tirosDe3 = 0; intentosDe1 = 0; intentosDe2 = 0; intentosDe3 = 0
    porcentajeDe1 = 0; porcentajeDe2 = 0; porcentajeDe3 = 0
    
    for jug in jugsEqP:
        valoracion += jug.valoracion
        rebotesDef += jug.rebotesDefensivos
        rebotesOf += jug.rebotesOfensivos
        robos += jug.robos
        perdidas += jug.perdidas
        asistencias += jug.asistencias
        tapones += jug.tapones
        taponesRecibidos += jug.taponesRecibidos
        faltas += jug.faltas
        faltasRecibidas += jug.faltasRecibidas
        contraataques += jug.contraataques
        tirosDe1 += jug.tirosDe1
        tirosDe2 += jug.tirosDe2
        tirosDe3 += jug.tirosDe3
        intentosDe1 += jug.intentosDe1
        intentosDe2 += jug.intentosDe2
        intentosDe3 += jug.intentosDe3
    porcentajeDe1 = tirosDe1 / intentosDe1 * 100
    porcentajeDe2 = tirosDe2 / intentosDe2 * 100
    porcentajeDe3 = tirosDe3 / intentosDe3 * 100

    dictEqP.update({"REBOTES_OF_P":rebotesOf, "REBOTES_DEF_P":rebotesDef, "FALTAS_P":faltas,
    "FALTAS_RECIBIDAS_P":faltasRecibidas, "TAPONES_P":tapones, "TAPONES_RECIBIDOS_P":taponesRecibidos,
    "CONTRAATAQUES_P":contraataques, "ASISTENCIAS_P":asistencias, "PERDIDAS_P":perdidas,"VALORACION_P":valoracion,
    "ROBOS_P":robos, "TL_TIRADOS_P":intentosDe1, "TL_ANOTADOS_P":tirosDe1, "PORCENTAJE_TL_P":porcentajeDe1,
    "T2_TIRADOS_P":intentosDe2, "T2_ANOTADOS_P":tirosDe2, "PORCENTAJE_T2_P":porcentajeDe2,
    "T3_TIRADOS_P":intentosDe3, "T3_ANOTADOS_P":tirosDe3, "PORCENTAJE_T3_P":porcentajeDe3})
    #TODO: Recuperar árbitros y entrenadores
